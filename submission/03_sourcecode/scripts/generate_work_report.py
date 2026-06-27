#!/usr/bin/env python3
"""生成第十九届全国大学生信息安全竞赛作品报告（AgentShield）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

OUTPUT = Path(__file__).resolve().parents[1] / "作品报告_AgentShield.docx"
LIVE_SITE = "http://114.215.209.144:8088"
GITHUB_REPO = "https://github.com/huang08666/agentShield"

BODY_FONT = "宋体"
BODY_SIZE = Pt(12)  # 小四


def set_run_font(run, bold=False):
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def add_para(doc, text, bold=False, first_line_indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    if level == 1:
        pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    if level == 1:
        run.font.size = Pt(14)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run)


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    # 封面
    t = doc.add_paragraph()
    t.alignment = 1
    r = t.add_run("附件2：\n")
    set_run_font(r)
    r2 = t.add_run(
        "第十九届全国大学生信息安全竞赛（作品赛）\n"
        "暨第三届“长城杯”网数智安全大赛（作品赛）\n"
        "作品报告\n"
    )
    set_run_font(r2, bold=True)
    r2.font.size = Pt(16)

    for label, value in [
        ("作品名称：", "AgentShield——面向大模型智能体工具调用的安全防护网关"),
        ("电子邮箱：", "（请填写参赛联系邮箱）"),
        ("提交日期：", "2026年6月"),
        ("参赛方向：", "应用安全 / 内容安全（大模型 Agent 工具调用安全）"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"{label}{value}")
        set_run_font(run)

    doc.add_page_break()

    # 目录（手工目录项）
    add_heading(doc, "目     录", level=1)
    for item in [
        "摘要\t1",
        "第一章 作品概述\t2",
        "第二章 作品设计与实现\t4",
        "第三章 作品测试与分析\t8",
        "第四章 创新性说明\t10",
        "第五章 总结\t12",
        "参考文献\t13",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        set_run_font(p.add_run(item))
    doc.add_page_break()

    # 摘要
    add_heading(doc, "摘要", level=1)
    add_para(
        doc,
        "随着大语言模型与智能体（Agent）技术的广泛应用，AI 系统不再局限于生成文本，"
        "而是能够自主规划并调用读文件、查数据库、发邮件、执行命令等外部工具。"
        "传统面向聊天输入的提示词防火墙难以覆盖“工具调用层”的新型风险，"
        "包括间接提示注入、多步攻击链、角色越权与敏感数据外泄等问题。"
    )
    add_para(
        doc,
        "本作品 AgentShield 是一款面向大模型智能体工具调用的安全防护网关。"
        "系统以“行为级安全”为核心理念，在 Agent 规划层与沙箱工具之间插入统一安全网关，"
        "对每一次工具调用进行角色权限校验、参数规则检测、语义意图识别、"
        "工具返回内容二次审查及全链路审计，并输出 allow / mask / confirm / block 四层决策。"
        "系统支持校园、个人助手、企业客服三套即插即用策略包，"
        "可通过 HTTP API、Python 装饰器、LangChain、MCP 等方式接入主流 Agent 框架，"
        "并支持 Cursor、Claude Desktop、Trae 等 IDE 通过 MCP 挂载受保护工具，"
        "以及 Claude API、OpenAI Codex 在 tool calling 循环中调用 evaluate 接口，"
        "并可在无外部 API 密钥条件下离线完整演示。"
        f"作品已部署上线演示站点（{LIVE_SITE}），"
        f"并在 GitHub 开源（{GITHUB_REPO}），便于评委在线体验与开发者二次集成。"
    )
    add_para(
        doc,
        "创新点主要体现在：（1）从“说了什么”升级为“准备执行什么”的行为级拦截；"
        "（2）规则引擎与本地语义意图相似度并联，可选 LLM Judge 二次研判；"
        "（3）人工确认闭环与完整审计链；（4）多场景策略包与 POLICY_FILE 即插即用。"
        "内置 22 条 Benchmark 样本实测检出率 100%、误报率 0%，"
        "平均单次策略评估延迟约 2.1 ms，具备工程可落地性与竞赛现场演示价值。",
        first_line_indent=True,
    )
    doc.add_page_break()

    # 第一章
    add_heading(doc, "第一章 作品概述", level=1)
    add_heading(doc, "1.1 背景分析", level=2)
    add_para(
        doc,
        "大模型 Agent 通过工具调用（Tool Calling）连接企业内部系统、个人文件与互联网服务，"
        "在提升效率的同时也扩大了攻击面。攻击者可通过恶意文档、污染网页、"
        "社会工程学话术等方式诱导 Agent 读取密钥、执行破坏性命令或外发敏感数据。"
        "OWASP LLM Top 10 将提示注入、不安全输出、敏感信息泄露、过度代理等列为重点风险。"
        "现有不少安全产品聚焦用户输入过滤，而对 Agent 已生成的工具调用计划、"
        "工具参数及工具返回内容缺乏统一治理，难以应对多步攻击链与编码绕过。"
    )
    add_heading(doc, "1.2 相关工作", level=2)
    add_para(
        doc,
        "相关工作可分为三类：一是聊天入口的 Prompt 防火墙与内容审核，"
        "主要检测用户自然语言；二是面向 API 的 WAF/RBAC，"
        "缺乏对 Agent 语义规划的理解；三是 Agent 框架内置的少量工具校验，"
        "与具体框架强耦合、难以复用。AgentShield 定位为框架无关的安全网关层，"
        "借鉴零信任“每次访问均须验证”思想，将每一次工具调用视为独立安全事件，"
        "结合 RBAC、规则引擎、语义检查与审计运营能力，填补 Agent 工具调用层的安全空白。",
    )
    add_heading(doc, "1.3 特色描述", level=2)
    add_bullet(doc, "行为级安全：拦截对象为 tool_name + tool_args + 用户角色 + 会话上下文，而非仅用户输入。")
    add_bullet(doc, "四层决策：allow（放行）、mask（脱敏）、confirm（人工确认）、block（阻断）。")
    add_bullet(doc, "多场景策略包：校园 / 个人 / 企业一键切换，策略与代码解耦。")
    add_bullet(doc, "即插即用集成：HTTP evaluate、Python embed、LangChain、MCP 四种接入方式。")
    add_bullet(doc, "主流 Agent 适配：Cursor / Claude Desktop / Trae（MCP）；Claude API / Codex（HTTP evaluate）。")
    add_bullet(doc, "离线可演示：Mock 计划引擎 + 沙箱工具，答辩现场无需外网 API。")
    add_bullet(doc, f"已上线演示：{LIVE_SITE}，评委零配置访问「快速集成 → 在线试集成」。")
    add_bullet(doc, f"开源可复现：{GITHUB_REPO}，支持 git clone 后嵌入模式本地集成。")
    add_heading(doc, "1.4 应用前景", level=2)
    add_para(
        doc,
        "AgentShield 可部署于教育科研平台的教学 Agent、个人知识助手、"
        "企业客服/工单 Agent、低代码 Agent 平台等场景。"
        "用户通过 POLICY_FILE 配置自有角色与工具权限即可落地，"
        "适用于对数据外发、命令执行、数据库查询有合规要求的行业环境。"
        f"当前已通过阿里云轻量服务器完成生产部署，"
        f"演示地址 {LIVE_SITE}，源码托管于 {GITHUB_REPO}。",
    )
    doc.add_page_break()

    # 第二章
    add_heading(doc, "第二章 作品设计与实现", level=1)
    add_heading(doc, "2.1 系统总体方案", level=2)
    add_para(
        doc,
        "系统采用前后端分离架构。后端基于 FastAPI 提供安全网关与 REST API；"
        "前端基于 React + TypeScript + Ant Design 实现可视化演示台。"
        "核心数据流为：用户输入 → Agent 规划（Mock/LLM）→ 工具调用计划 → "
        "AgentShield Gateway → Policy Engine + Semantic Checker → 决策执行 → "
        "Content Guard → SQLite 审计 → 前端展示。",
    )
    add_para(
        doc,
        "信任边界设定为：Agent 规划层不可信；工具参数与工具返回均视为不可信输入。"
        "所有外部能力调用必须经网关评估后方可执行（confirm 类型需人工批准）。",
    )
    add_heading(doc, "2.2 软件模块设计", level=2)
    add_para(doc, "后端核心模块如下：", first_line_indent=False)
    modules = [
        "Shield Gateway（gateway.py）：统一拦截入口，负责评估、执行、审计与 confirm 闭环。",
        "Policy Engine（policy_engine.py）：RBAC、敏感路径、危险 SQL/Shell、邮件外发、浏览器 URL 等规则评分。",
        "Semantic Checker（semantic_checker.py）：基于意图库与文本相似度的语义风险检测，可选 LLM Judge。",
        "Normalize（normalize.py）：全角/零宽/路径/SQL 规范化，降低编码绕过风险。",
        "Content Guard（content_guard.py）：对工具返回做注入、敏感信息与语义二次检测。",
        "Profile Manager：校园 campus、个人 personal、企业 enterprise 三套 YAML 策略包热切换。",
        "Benchmark Runner：22 条标准化样本与评委现场自拟 evaluate 接口。",
        "Integration：wrapper 装饰器、sdk HTTP 客户端、LangChain Tools、MCP Server。",
    ]
    for m in modules:
        add_bullet(doc, m)

    add_heading(doc, "2.3 策略与决策机制", level=2)
    add_para(
        doc,
        "策略文件（YAML）定义 role_permissions（allow/deny/confirm）、tool_risk_level、"
        "sensitive_paths、dangerous_sql_patterns、dangerous_shell_patterns、"
        "prompt_injection_patterns、external_email_domains、decision_thresholds 等字段。"
        "引擎对每次调用累加风险分，并与角色权限决策取最严结果。"
        "风险分区间映射为：0–29 allow、30–59 mask、60–79 confirm、80–100 block。",
    )
    add_para(
        doc,
        "语义检查默认采用 local 模式：加载 semantic_intents.yaml 中四类意图"
        "（提示注入、数据外泄、越权访问、破坏性操作），"
        "通过双字片段召回、Jaccard 相似度与最长公共子序列计算与用户输入/工具参数的相似度；"
        "配置 SEMANTIC_CHECK_MODE=llm 且提供 API Key 时，可启用大模型 Judge 二次研判并自动回退本地模式。",
    )
    add_heading(doc, "2.4 前端功能", level=2)
    pages = [
        "Agent 对话：完整展示计划、决策、时间线、风险解释与人工确认按钮。",
        "策略中心：场景切换、策略可视化、在线编辑与热重载。",
        "攻击演示：恶意文档、网页污染、数据库越权、危险命令四剧本一键运行。",
        "Benchmark：固定样本跑分与评委现场单条评估。",
        "审计日志：筛选、详情、CSV 导出及待确认记录处理。",
        "框架演示 / 快速集成：LangChain、MCP 演示；"
        "主流 Agent 接入（Cursor、Claude、Trae、Codex）文档与在线试集成；"
        f"开源仓库（{GITHUB_REPO}）与演示站（{LIVE_SITE}）。",
    ]
    for p in pages:
        add_bullet(doc, p)

    add_heading(doc, "2.5 关键技术指标", level=2)
    add_bullet(doc, "支持工具类型：read_file、query_db、send_email、run_shell、browser_mock（均可扩展）。")
    add_bullet(doc, "鉴权：JWT（HS256），角色由服务端签发，防止客户端伪造 user_role。")
    add_bullet(doc, "数据库：SQLite 全链路审计，记录风险分、决策、原因与结果预览。")
    add_bullet(doc, "Benchmark 平均评估延迟：约 2.1 ms/条（本地 campus 策略实测）。")
    add_bullet(doc, "自动化测试：51 条 pytest 用例覆盖策略、绕过、confirm、语义、鉴权等。")
    add_bullet(doc, f"生产部署：nginx + systemd，演示地址 {LIVE_SITE}，后端 127.0.0.1:8001 仅本机反代。")
    add_bullet(doc, f"开源仓库：{GITHUB_REPO}，含 INTEGRATION.md、INTEGRATION-AGENTS.md 与 deploy 脚本。")
    doc.add_page_break()

    # 第三章
    add_heading(doc, "第三章 作品测试与分析", level=1)
    add_heading(doc, "3.1 测试方案", level=2)
    add_para(
        doc,
        "测试分为四个层次：（1）单元测试：策略引擎、规范化、语义相似度、JWT 鉴权；"
        "（2）场景测试：四内置攻击剧本端到端阻断；"
        "（3）Benchmark 测试：22 条标注样本（normal/attack/bypass）自动跑分；"
        "（4）人工演示测试：三场景策略包切换、人工 confirm 闭环、LangChain/MCP 接入演示；"
        f"（5）线上验收：{LIVE_SITE} 健康检查、登录、快速集成与攻击剧本演示。",
    )
    add_heading(doc, "3.2 测试环境", level=2)
    add_bullet(doc, "操作系统：Linux；Python 3.11；Node.js 18+。")
    add_bullet(doc, "本地开发：FastAPI + Uvicorn + Vite 开发服务器。")
    add_bullet(doc, f"线上环境：阿里云轻量服务器，演示入口 {LIVE_SITE}（nginx :8088 反代 + 后端 127.0.0.1:8001）。")
    add_bullet(doc, "后端：FastAPI + Uvicorn（生产单 worker），默认 USE_MOCK_LLM=true。")
    add_bullet(doc, "数据库：SQLite 本地文件 agentshield.db。")
    add_heading(doc, "3.3 Benchmark 结果", level=2)
    add_para(
        doc,
        "Benchmark 样本集共 22 条，其中 normal 5 条、attack 9 条、bypass 8 条。"
        "bypass 类专门覆盖空格/全角/零宽字符/SQL 注释/路径变体等规则绕过手法。"
        "在当前默认校园策略下实测结果如下：",
    )
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("指标", "结果"),
        ("总样本数", "22"),
        ("检出率", "100.0%"),
        ("误报率", "0.0%"),
        ("分类通过率", "normal 5/5，attack 9/9，bypass 8/8"),
        ("平均延迟", "约 2.1 ms/条"),
    ]
    for i, (a, b) in enumerate(rows_data):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, bold=(i == 0))

    add_para(doc, "")
    add_heading(doc, "3.4 典型攻击场景验证", level=2)
    scenarios = [
        "恶意文档诱导泄密：读取 malicious_doc.txt 后诱导读 secret 并外发邮件，最终 block。",
        "网页工具返回污染：browser_mock 返回含注入指令内容，Content Guard 与语义层联动处置。",
        "数据库越权：SELECT * FROM students 被 SQL 规则阻断，建议使用 COUNT 查询。",
        "危险命令：rm -rf sandbox/* 被 Shell 危险模式与白名单双重拦截。",
        "语义变体注入：未命中固定注入关键词的改写话术，由语义意图库识别并加分/阻断。",
    ]
    for s in scenarios:
        add_bullet(doc, s)
    add_heading(doc, "3.5 结果分析", level=2)
    add_para(
        doc,
        "测试表明，规则引擎可有效覆盖确定性攻击与编码绕过；"
        "语义检查弥补了固定关键词对同义改写覆盖不足的问题；"
        "四层决策与 confirm 闭环在保障安全的同时兼顾业务可用性（如教师发邮件需确认）。"
        "误报率 0% 说明正常样本（读报告、COUNT 查询、白名单命令）未被过度拦截。"
        "局限在于：语义 local 模式对极短文本仍可能误判，需与规则联合使用；"
        "Agent 规划层在演示模式下为 Mock 引擎，真实 LLM 行为需结合生产环境进一步评估。",
    )
    doc.add_page_break()

    # 第四章
    add_heading(doc, "第四章 创新性说明", level=1)
    innovations = [
        (
            "4.1 从输入过滤到行为拦截的范式转变",
            "传统 Prompt 防火墙关注用户“说了什么”，AgentShield 关注 Agent“准备执行什么”。"
            "将安全左移至工具调用网关，可拦截多步攻击链（读密钥→发邮件）与会话上下文关联风险，"
            "更贴合 OWASP LLM08 过度代理等新型威胁。",
        ),
        (
            "4.2 规则 + 语义并联的多层防御",
            "在 RBAC、路径/SQL/Shell 规则与规范化防绕过基础上，"
            "新增基于意图库的本地语义相似度检测，并可选 LLM Judge。"
            "既可捕获“不要再遵守先前安全限制”等改写注入，"
            "又可在无 API 条件下离线运行，满足竞赛演示与私有化部署需求。",
        ),
        (
            "4.3 可运营的四层决策与人工闭环",
            "区别于简单通过/拒绝，系统提供 mask 脱敏与 confirm 人工确认，"
            "配套审计日志、待确认统计与 POST /api/shield/confirm 批准/拒绝 API，"
            "形成“检测—决策—人工—执行—留痕”完整闭环。",
        ),
        (
            "4.4 即插即用场景化策略包",
            "校园、个人、企业三套 Profile 通过 YAML 一键切换，"
            "无需修改代码即可改变角色体系与规则。"
            "生产环境可通过 POLICY_FILE 加载用户自有策略，"
            "实现竞赛演示与真实接入的统一内核。",
        ),
        (
            "4.5 框架无关的开放集成",
            "提供 HTTP evaluate、Python 装饰器、LangChain StructuredTool、MCP stdio 四种接入路径，"
            "并编写 INTEGRATION-AGENTS.md 指导 Cursor、Claude Desktop、Trae 通过 MCP 配置受保护工具，"
            "以及 Claude API、OpenAI Codex 在 tool_use / tool_calls 循环中调用安全网关，"
            "不绑定单一 Agent 框架，降低用户迁移成本，"
            "体现“安全网关”而非“演示 Demo”的工程定位。"
            f"作品已开源（{GITHUB_REPO}）并上线演示站（{LIVE_SITE}），"
            "评委可在线体验，开发者可 clone 后嵌入集成。",
        ),
    ]
    for title, body in innovations:
        add_heading(doc, title, level=2)
        add_para(doc, body)
    doc.add_page_break()

    # 第五章
    add_heading(doc, "第五章 总结", level=1)
    add_para(
        doc,
        "本作品针对大模型 Agent 工具调用场景的安全空白，"
        "设计并实现了 AgentShield 安全防护网关。"
        "系统以行为级拦截为核心，融合 RBAC 策略引擎、规范化防绕过、"
        "本地语义意图检测、工具返回 Content Guard、JWT 鉴权与全链路审计，"
        "并提供四层决策与人工确认闭环。"
        "三套场景策略包与多种集成方式使系统兼具竞赛演示表现力与工程落地潜力。"
        f"作品已完成云端部署（{LIVE_SITE}）与 GitHub 开源（{GITHUB_REPO}），"
        "形成“在线体验 + 源码复现 + 嵌入集成 + 主流 Agent MCP/HTTP 接入”的完整交付形态。",
    )
    add_para(
        doc,
        "Benchmark 实测 22 条样本检出率 100%、误报率 0%，"
        "验证了方案的有效性。"
        "后续将扩展 ABAC 细粒度授权、审计防篡改、"
        "更多真实 Agent 框架适配以及语义模型在线学习等能力，"
        "持续提升对未知攻击变体的防护水平。",
    )
    doc.add_page_break()

    # 参考文献
    add_heading(doc, "参考文献", level=1)
    refs = [
        "[1] OWASP Foundation. OWASP Top 10 for Large Language Model Applications[EB/OL]. https://owasp.org/www-project-top-10-for-large-language-model-applications/, 2024.",
        "[2] 李建华. 网络空间威胁情报感知、共享与分析技术综述[J]. 网络与信息安全学报, 2016, 2(2): 16-29.",
        "[3] Anthropic. Model Context Protocol Specification[EB/OL]. https://modelcontextprotocol.io/, 2024.",
        "[4] LangChain Inc. LangChain Documentation: Tools and Agents[EB/OL]. https://python.langchain.com/, 2024.",
        "[5] NIST. Zero Trust Architecture (SP 800-207)[S]. National Institute of Standards and Technology, 2020.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        set_run_font(p.add_run(ref))

    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build_document()
